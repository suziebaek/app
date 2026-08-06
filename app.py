import streamlit as st
import pandas as pd
from docx import Document
import pdfplumber
import io
import re
import zipfile
import datetime
from openpyxl.styles import PatternFill

# 1. 파일명 규칙 정제 함수

# vocabulary 및 발음 파일명 모두에 적용: sb/sth 계열 줄임말 풀어쓰기
# (sb/sth처럼 함께 붙어있는 경우를 먼저 처리해야 단독 sb, sth 규칙에 의해 잘못 치환되지 않음)
def expand_sb_sth(text):
    t = text
    t = re.sub(r'\bsb\s*/\s*sth\b', 'something', t)   # sb/sth -> something
    t = re.sub(r'\bsb\b', 'somebody', t)              # sb -> somebody
    t = re.sub(r'\bsth\b', 'something', t)            # sth -> something
    return t

# 발음 파일명에서만 추가로 사용하지 않을 줄임말 -> 원래 표현 매핑
def expand_abbreviations(text):
    t = expand_sb_sth(text)
    t = re.sub(r'\bV-ing\b', 'ing', t)                # V-ing -> ing
    t = re.sub(r'\bto\s+V\b', 'to', t)                # to V -> to
    return t

def clean_filename(text):
    text = expand_abbreviations(text)
    text = text.replace(" ", "_")
    text = re.sub(r'[^a-zA-Z0-9_]', '', text)
    # 대소문자는 vocabulary 원문 그대로 유지 (예: protect A from B -> protect_A_from_B.mp3)
    return text + ".mp3"

# 챕터/회차 표기 감지: "Chapter 1", "CHAPTER 01", "CH01", "CH02", "1회", "04회" 등 지원
CHAPTER_LINE_PATTERN = re.compile(r'(?i)^(?:chapter|ch)\.?\s*0*(\d+)\b')
ROUND_LINE_PATTERN = re.compile(r'^제?\s*0*(\d+)\s*회\s*$')

def detect_chapter_number(line):
    stripped = line.strip()
    m = CHAPTER_LINE_PATTERN.match(stripped)
    if m:
        return int(m.group(1))
    m = ROUND_LINE_PATTERN.match(stripped)
    if m:
        return int(m.group(1))
    return None

def is_round_marker(line):
    return ROUND_LINE_PATTERN.match(line.strip()) is not None

# 2. 워드 파싱 함수
def parse_word_content(docx_file):
    doc = Document(docx_file)
    data = []
    raw_texts = []
    
    current_chapter_num = 1
    current_week_title = "Week01/"
    current_week_num = 1
    
    paragraphs = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        # Shift+Enter(줄바꿈)로 한 문단 안에 여러 줄(예: "word: 뜻" 다음 줄에 "Syn: ...")이
        # 들어있는 경우를 대비해 문단 내부의 줄바꿈 기준으로도 분리
        for sub in p.text.split("\n"):
            sub = sub.strip()
            if sub:
                paragraphs.append(sub)

    # 번호는 "1. word" (마침표)와 "01  word" (공백 2칸) 형식을 모두 지원
    entry_pattern = re.compile(r'^(?:(\d+)[\.\s]+)?([^:]+):\s*(.*)$')
    keyword_only_pattern = re.compile(r'(?i)^(Syn|Ant|Atn|유의어|반의어)$')
    tail_keyword_pattern = re.compile(r'(?i)^(Syn|Ant|Atn|유의어|반의어)\s*:\s*(.*)$')

    # "1회", "04회"처럼 "~회" 표기가 문서 어딘가에 하나라도 있으면, 그 표기가
    # 누락된 곳도 있을 수 있다고 보고 번호가 정확히 1로 리셋되는 지점(예: ...20. -> 1.)을
    # 새 회차 시작으로 자동 보완합니다. "CH01" 같은 챕터 표기만 쓰는 문서(예: 회차 내부에
    # 번호가 우연히 반복/오탈자로 흐트러지는 경우)에는 적용하지 않아 기존 동작을 유지합니다.
    use_round_reset_fallback = any(is_round_marker(l) for l in paragraphs)
    previous_entry_num = None
    just_saw_explicit_marker = False

    for line in paragraphs:
        raw_texts.append(line)
        
        # 챕터 감지 (Week 번호 연동)
        chapter_num = detect_chapter_number(line)
        if chapter_num is not None:
            current_week_num = chapter_num
            current_week_title = f"Week{current_week_num:02d}/"
            current_chapter_num = 1
            previous_entry_num = None
            just_saw_explicit_marker = True
            continue

        # 단어: 뜻 [탭/2칸 이상 공백 + Syn/Ant: 유의어] 추출
        match = entry_pattern.match(line)

        if match:
            entry_num_str = match.group(1)
            raw_word = expand_sb_sth(match.group(2).strip())

            # "Syn: xxx" 처럼 단어 없이 유의어만 있는 줄은 아래 유의어 처리 블록으로 넘김
            if not keyword_only_pattern.match(raw_word):
                # 번호가 이전보다 작아지며 리셋되면 새 회차(챕터) 시작으로 간주
                # (단, 직전에 명시적 챕터/회차 마커를 이미 처리했다면 중복 증가하지 않음)
                if entry_num_str is not None:
                    entry_num = int(entry_num_str)
                    # 번호가 정확히 1로 리셋되는 경우만 새 회차 시작으로 인정
                    # (오탈자로 인한 불규칙한 감소는 무시하여 오탐을 방지)
                    if (use_round_reset_fallback and previous_entry_num is not None
                            and entry_num == 1 and previous_entry_num != 1
                            and not just_saw_explicit_marker):
                        current_week_num += 1
                        current_week_title = f"Week{current_week_num:02d}/"
                        current_chapter_num = 1
                    previous_entry_num = entry_num

                just_saw_explicit_marker = False

                rest = match.group(3)

                # 뜻과 유의어/반의어를 탭 또는 2칸 이상 공백으로 분리
                parts = re.split(r'\t|\s{2,}', rest, maxsplit=1)
                meaning = parts[0].strip()

                extra_info = ""
                if len(parts) > 1:
                    tail = parts[1].strip()
                    tail_match = tail_keyword_pattern.match(tail)
                    if tail_match:
                        extra_info = tail_match.group(2).strip()

                if re.search(r'[가-힣]', meaning):
                    sound_file = clean_filename(raw_word)

                    data.append({
                        "lesson_title": current_week_title,
                        "lesson_order_seq": current_week_num,
                        "page_order_seq": current_chapter_num,
                        "vocabulary": raw_word,
                        "vocabulary_kor": meaning,
                        "vocabulary_sound": sound_file,
                        "vocabulary_excep": extra_info,
                        "prompt": ""
                    })
                    current_chapter_num += 1
                continue
        
        # 유의어 처리 (Syn/Ant가 별도의 줄로 존재하는 예전 형식 지원)
        if any(keyword in line for keyword in ["Syn", "Atn", "Ant", "유의어", "반의어"]):
            extra_info = line.split(":", 1)[-1].strip()
            if data:
                first_extra = re.split(r'\s{2,}', extra_info)[0].strip()
                data[-1]["vocabulary_excep"] = first_extra

    return data, raw_texts

# 2-0. 워드 파싱 함수 (용어집/Glossary 형식)
# "word / 뜻 / (부가 설명) / syn: ... / ant: ... / 예문 / 활용 표현" 처럼
# 각 항목이 한 줄에 한 요소씩, 콜론으로 단어-뜻이 묶이지 않고 여러 줄에 나뉘어
# 있는 문서 형식을 지원합니다. 단어 줄은 "그 줄 자체엔 한글이 없고, 바로 다음
# 줄에는 한글이 있다"는 특징으로 판별합니다.
def _has_korean(text):
    return bool(re.search(r'[가-힣]', text))

def _is_glossary_word_start(lines, idx):
    if idx >= len(lines):
        return False
    line = lines[idx]
    if re.match(r'(?i)^(syn|ant|atn|유의어|반의어)\s*:', line):
        return False
    if _has_korean(line):
        return False
    if detect_chapter_number(line) is not None:
        return False
    if idx + 1 >= len(lines):
        return False
    return _has_korean(lines[idx + 1])

def parse_word_content_glossary_style(docx_file):
    doc = Document(docx_file)

    # 워드 안에서 줄바꿈(Shift+Enter)으로 나뉜 줄까지 모두 개별 라인으로 펼침
    raw_lines = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for sub in p.text.split("\n"):
            sub = sub.strip()
            if sub:
                raw_lines.append(sub)

    data = []
    raw_texts = list(raw_lines)

    current_chapter_num = 1
    current_week_title = "Week01/"
    current_week_num = 1

    tail_keyword_pattern = re.compile(r'(?i)^(syn|ant|atn|유의어|반의어)\s*:\s*(.*)$')

    i = 0
    n = len(raw_lines)
    while i < n:
        line = raw_lines[i]

        # 챕터 감지 (Week 번호 연동)
        chapter_num = detect_chapter_number(line)
        if chapter_num is not None:
            current_week_num = chapter_num
            current_week_title = f"Week{current_week_num:02d}/"
            current_chapter_num = 1
            i += 1
            continue

        if _is_glossary_word_start(raw_lines, i):
            raw_word = expand_sb_sth(line)
            meaning = raw_lines[i + 1].strip()

            # 단어 다음, 다음 단어(혹은 챕터)가 나오기 전까지의 구간에서
            # 첫 번째로 등장하는 syn/ant 값을 유의어/반의어로 채택
            extra_info = ""
            j = i + 2
            while j < n and not _is_glossary_word_start(raw_lines, j) and detect_chapter_number(raw_lines[j]) is None:
                tail_match = tail_keyword_pattern.match(raw_lines[j])
                if tail_match and not extra_info:
                    extra_info = tail_match.group(2).strip()
                j += 1

            if _has_korean(meaning):
                sound_file = clean_filename(raw_word)
                data.append({
                    "lesson_title": current_week_title,
                    "lesson_order_seq": current_week_num,
                    "page_order_seq": current_chapter_num,
                    "vocabulary": raw_word,
                    "vocabulary_kor": meaning,
                    "vocabulary_sound": sound_file,
                    "vocabulary_excep": extra_info,
                    "prompt": ""
                })
                current_chapter_num += 1

            i = j
            continue

        i += 1

    return data, raw_texts

# 2-1. PDF 파싱 함수
def parse_pdf_content(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        full_text = "\n".join(page.extract_text() or "" for page in pdf.pages)

    lines = [ln.strip() for ln in full_text.split("\n") if ln.strip()]

    data = []
    raw_texts = []

    current_chapter_num = 1
    current_week_title = "Week01/"
    current_week_num = 1

    # PDF 텍스트 추출 시 "01 word: 뜻   Syn: xxx" 처럼 번호가 마침표 없이
    # 붙는 형태와, 뜻과 Syn/Ant 사이가 한 칸 공백으로만 구분되는 경우가 많아
    # 워드용 정규식과는 별도의 패턴을 사용합니다.
    entry_pattern = re.compile(r'^\d{1,2}\s+([^:]+):\s*(.*)$')
    syn_pattern = re.compile(r'\b(Syn|Ant)\s*:\s*(.+)$')

    for line in lines:
        raw_texts.append(line)

        # 챕터 감지 (Week 번호 연동)
        chapter_num = detect_chapter_number(line)
        if chapter_num is not None:
            current_week_num = chapter_num
            current_week_title = f"Week{current_week_num:02d}/"
            current_chapter_num = 1
            continue

        # 번호 + 단어: 뜻 [Syn/Ant: 유의어] 추출
        match = entry_pattern.match(line)
        if match:
            raw_word = expand_sb_sth(match.group(1).strip())
            rest = match.group(2).strip()

            syn_match = syn_pattern.search(rest)
            if syn_match:
                meaning = rest[:syn_match.start()].strip()
                extra_info = syn_match.group(2).strip()
            else:
                meaning = rest
                extra_info = ""

            if re.search(r'[가-힣]', meaning):
                sound_file = clean_filename(raw_word)

                data.append({
                    "lesson_title": current_week_title,
                    "lesson_order_seq": current_week_num,
                    "page_order_seq": current_chapter_num,
                    "vocabulary": raw_word,
                    "vocabulary_kor": meaning,
                    "vocabulary_sound": sound_file,
                    "vocabulary_excep": extra_info,
                    "prompt": ""
                })
                current_chapter_num += 1

    return data, raw_texts

# 4. 엑셀 생성 헬퍼 (지정된 컬럼에 회색 배경 적용)
def build_excel_bytes(df_subset):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_subset.to_excel(writer, index=False, sheet_name='원고작성')
        worksheet = writer.sheets['원고작성']

        # RGB(191, 191, 191) 색상 (HEX: BFBFBF)
        gray_fill = PatternFill(start_color='BFBFBF', end_color='BFBFBF', fill_type='solid')

        # 색상 적용 대상 컬럼
        target_cols = ["service_code", "track_code", "top_cors_id", "component_code", "book_code", "act_code"]

        for col_name in target_cols:
            if col_name in df_subset.columns:
                col_idx = df_subset.columns.get_loc(col_name) + 1
                for row in range(1, len(df_subset) + 2):
                    cell = worksheet.cell(row=row, column=col_idx)
                    cell.fill = gray_fill

    return output.getvalue()

# --- 3. Streamlit UI (여기서 변수가 정의됩니다) ---
st.set_page_config(page_title="최종 단어 변환기", layout="wide")
st.title("📑 주차 연동 및 색상 지정 시스템")

# 사이드바 설정
with st.sidebar:
    st.header("⚙️ 고정값 설정")
    service_code = st.text_input("service_code", "SVC170")
    track_code = st.text_input("track_code", "RSV_TRK01")
    top_cors_id = st.text_input("top_cors_id", "1879")
    level_code = st.text_input("level_code", "TO_R_E_SP")
    component_code = st.text_input("component_code", "COM170")
    book_code = st.text_input("book_code", "SVC170")
    act_name = st.text_input("act_name", "Vocablist")
    show_debug = st.checkbox("디버그 모드", value=False)

    st.header("📁 파일명 설정")
    st.caption("예: TOT_AU_Recap_S_W04_260731.xlsx")
    file_name_prefix = st.text_input("파일명 접두어 (주차/날짜 앞부분)", "TOT_AU_Recap_S")
    file_date = st.date_input("파일명 날짜", value=datetime.date.today())
    date_str = file_date.strftime("%y%m%d")

# 중요: 여기서 uploaded_file 변수가 정의됩니다!
uploaded_file = st.file_uploader("워드(.docx) 또는 PDF(.pdf) 파일을 업로드하세요", type=["docx", "pdf"])

# 변수가 정의된 이후에 사용합니다.
if uploaded_file is not None:
    file_ext = uploaded_file.name.lower().rsplit(".", 1)[-1]

    if file_ext == "pdf":
        parsed_data, raw_texts = parse_pdf_content(uploaded_file)
    else:
        parsed_data, raw_texts = parse_word_content(uploaded_file)
        # 기본 형식(단어: 뜻)으로 하나도 못 찾으면, 용어집(Glossary) 형식으로 재시도
        if not parsed_data:
            uploaded_file.seek(0)
            parsed_data, raw_texts = parse_word_content_glossary_style(uploaded_file)
    
    if parsed_data:
        df = pd.DataFrame(parsed_data)
        
        # 고정값 채우기
        df['service_code'] = service_code
        df['track_code'] = track_code
        df['top_cors_id'] = top_cors_id
        df['level_code'] = level_code
        df['component_code'] = component_code
        df['book_code'] = book_code
        df['act_code'] = "RSV_ACT002"
        df['act_name'] = act_name
        
        # 컬럼 순서 재배치
        final_cols = [
            "service_code", "track_code", "top_cors_id", "level_code", 
            "component_code", "book_code", "lesson_order_seq", "lesson_title", 
            "act_code", "act_name", "page_order_seq", "vocabulary", 
            "vocabulary_kor", "vocabulary_sound", "vocabulary_excep", "prompt"
        ]
        df_final = df.reindex(columns=final_cols).fillna("")

        st.success("✅ 분석 완료!")
        st.dataframe(df_final)

        # 전체 결과물(모든 주차 통합) 엑셀
        excel_bytes_all = build_excel_bytes(df_final)
        all_file_name = f"{file_name_prefix}_{date_str}.xlsx"

        st.download_button(
            label="📊 전체 결과물 다운로드 (모든 주차 통합)",
            data=excel_bytes_all,
            file_name=all_file_name,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        # 주차별 파일 생성 (개별 다운로드 + ZIP 일괄 다운로드)
        st.subheader("📅 주차별 다운로드")

        week_numbers = sorted(df_final["lesson_order_seq"].unique())
        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for week_num in week_numbers:
                week_df = df_final[df_final["lesson_order_seq"] == week_num]
                week_bytes = build_excel_bytes(week_df)
                week_file_name = f"{file_name_prefix}_W{int(week_num):02d}_{date_str}.xlsx"
                zf.writestr(week_file_name, week_bytes)

                st.download_button(
                    label=f"⬇️ Week {int(week_num):02d} 다운로드 ({len(week_df)}개 단어)",
                    data=week_bytes,
                    file_name=week_file_name,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key=f"week_download_{week_num}"
                )

        st.download_button(
            label="🗂️ 주차별 파일 전체 ZIP으로 한꺼번에 다운로드",
            data=zip_buffer.getvalue(),
            file_name=f"{file_name_prefix}_ALL_{date_str}.zip",
            mime="application/zip"
        )
    
    if show_debug:
        with st.expander("🔍 디버그 데이터"):
            for i, txt in enumerate(raw_texts):
                st.text(f"L{i+1}: {txt}")
